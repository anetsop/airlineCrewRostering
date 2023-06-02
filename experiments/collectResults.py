import openpyxl
import docx
import os

def collectCSOResults(folderName, fileName):
    # collect results from multiple excel files 
    # created from runs using multi-step CSO
    # and combine them into a single table of a
    # microsoft word document
    folder = os.scandir(folderName)
    files = [entry for entry in folder if entry.is_file() == True]
    aggregate = []
    document = docx.Document()
    table = document.add_table(rows=1, cols=8)
    row = table.rows[0].cells
    row[0].text = "Execution"
    row[1].text = "FL"
    row[2].text = "Best"
    row[3].text = "Worst"
    row[4].text = "Sols"
    row[5].text = "Paths"
    row[6].text = "Similarity"
    row[7].text = "Jumps"
    for file in files:
        workbook = openpyxl.load_workbook(filename = file.path)
        sheet1 = workbook["General Information"]
        sheet2 = workbook["Solution Statistics"]
        sheet3 = workbook["Optimization Algorithm"]
        # exec. time, params, best, worst, valid, total, jumps, similarity
        statList = (sheet1['D7'].value, sheet1['N8'].value, sheet2['D9'].value, sheet2['D10'].value, \
                    sheet3['D5'].value, sheet3['D7'].value, sheet3['D9'].value, round(sheet3['D10'].value*100, 2))
        aggregate.append(statList)
        workbook.close()
    for result in aggregate:
        print(result)
        row = table.add_row().cells
        row[0].text = str(result[0])
        row[1].text = str(result[1])
        row[2].text = str(result[2])
        row[3].text = str(result[3])
        row[4].text = str(result[4])
        row[5].text = str(result[5])
        row[6].text = str(result[7])
        row[7].text = str(result[6])
    document.save(fileName)

def collectAOAResults(folderName, fileName):
    # collect results from multiple excel files 
    # created from runs using AOA
    # and combine them into a single table of a
    # microsoft word document
    folder = os.scandir(folderName)
    files = [entry for entry in folder if entry.is_file() == True]
    aggregate = []
    document = docx.Document()
    table = document.add_table(rows=1, cols=11)
    row = table.rows[0].cells
    row[0].text = "Execution"
    row[1].text = "C1"
    row[2].text = "C2"
    row[3].text = "C3"
    row[4].text = "C4"
    row[5].text = "Best"
    row[6].text = "Worst"
    row[7].text = "Sols"
    row[8].text = "Paths"
    row[9].text = "Similarity"
    row[10].text = "Jumps"
    for file in files:
        workbook = openpyxl.load_workbook(filename = file.path)
        sheet1 = workbook["General Information"]
        sheet2 = workbook["Solution Statistics"]
        sheet3 = workbook["Optimization Algorithm"]
        # exec. time, params, best, worst, valid, total, jumps, similarity
        statList = (sheet1['D7'].value, sheet1['N8'].value, sheet1['N9'].value, sheet1['N10'].value, sheet1['N11'].value, \
                    sheet2['D9'].value, sheet2['D10'].value, \
                    sheet3['D5'].value, sheet3['D7'].value, sheet3['D9'].value, round(sheet3['D10'].value*100, 2))
        aggregate.append(statList)
        workbook.close()
    for result in aggregate:
        print(result)
        row = table.add_row().cells
        row[0].text = str(result[0])
        row[1].text = str(result[1])
        row[2].text = str(result[2])
        row[3].text = str(result[3])
        row[4].text = str(result[4])
        row[5].text = str(result[5])
        row[6].text = str(result[6])
        row[7].text = str(result[7])
        row[8].text = str(result[8])
        row[9].text = str(result[10])
        row[10].text = str(result[9])
    document.save(fileName)

if __name__ == "__main__":

    # collect the results for each seed used by AOAExperiment.py 
    # and csoExperiment.py. A word document is created for each seed
    # to collect and store its results.
    # For each seed we have already created a folder that contains
    # two subfolders "CSO" and "AOA". "CSO" folders contains the 
    # excel files created by running the app for the specific seed
    # using multi-step CSO and "AOA" folder contains the files
    # created running the app for the seed using AOA algorithm
    # For example, the path to the folder for the seed "1528461486438309900"
    # looks like this:
    # "PATH-TO-airline-crew-rostering\output\seed 1528461486438309900"

    suffixes = ["1528461486438309900", "1522171910314035713", \
                "1528661486938309900", "1622631426338403900", \
                "188989346674497446", "186215413753529927", \
                "1093348801906287046", "4850465083999709955", \
                "1887491532428972499", "1657125590155585700", \
                "1883334465705964489", "1515819194504058851", \
                "1715047472822920876", "1975167106717073911", \
                "1462034985755456096", "59236085", 
                "504240966", \
                "1190914300682851677", "85773837", \
                "50997379", "46343184", \
                "2446767", "25477884", \
                "759251", "1849292547409738640", \
                "1158728481045126211", "1630465166841426413", \
                "1475027173915230540", "1315435446491772171", \
                "1370799939199725866", "1243453791151397408", \
                "1404351686654699285", "1717925693767016444", \
                "1221673221631747003", \
                "1983251052612723770", "1353686709715744173", \
                "1841636645489128699", "1474082555734599704", \
                "1275993964133764308", "1964018949345832996", \
                "1901613396743436965", "1029409762021639785", \
                "1442291518289717303", "1290776829459136181"
                ]
    
    directory = os.getcwd()
    par_dir = os.path.dirname(directory)
    directory = par_dir + "\\airline_crew_rostering\\output\\seed "
    
    for suffix in suffixes:
        path1 = directory + suffix + "\\CSO"
        path2 = directory + + suffix + "\\AOA"
        file1 = "CSO_seed_"+ suffix + ".docx"
        file2 = "AOA_seed_"+ suffix + ".docx"
        collectCSOResults(path1, file1)
        collectAOAResults(path2, file2)
